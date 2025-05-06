import os
import sys
import shutil
import re
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx2pdf import convert

def center_docx_text(source_path, target_path=None):
    """
    Create a copy of a DOCX file and center all text in the new file.
    
    Args:
        source_path (str): Path to the source DOCX file
        target_path (str, optional): Path where the new DOCX file will be saved.
                                    If not provided, will use source filename with '_new' suffix.
    
    Returns:
        str: Path to the new new DOCX file, or empty string if operation failed
    """
    try:
        # Validate source file exists
        if not os.path.exists(source_path):
            print(f"Error: Source file not found at {source_path}")
            return ""
        
        # Create target path if not provided
        if not target_path:
            # Create a 'new_docs' folder in the same directory as the script
            script_dir = os.path.dirname(os.path.abspath(__file__))
            new_folder = os.path.join(script_dir, "new_docs")
            
            # Create the folder if it doesn't exist
            if not os.path.exists(new_folder):
                os.makedirs(new_folder)
                print(f"Created new folder: {new_folder}")
            
            file_name = os.path.basename(source_path)
            name, ext = os.path.splitext(file_name)
            target_path = os.path.join(new_folder, f"{name}_new{ext}")
        
        # Create a copy of the original file
        shutil.copy2(source_path, target_path)
        
        # Open the copied document
        doc = Document(target_path)
        
        # Process all paragraphs in the document
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():  # Only process non-empty paragraphs
                # Center the paragraph and set right-to-left text direction
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                
                # Set right-to-left text direction for the paragraph
                # Note: We can't use paragraph.paragraph_format.bidi as it's not available in python-docx
                # Instead, we'll rely on setting RTL at the run level
                
                # Set RTL for all existing runs
                for run in paragraph.runs:
                    run.font.rtl = True
                
                # We'll process each run separately to preserve formatting
                if any(char.isdigit() for char in paragraph.text):
                    # Store the original runs with their formatting
                    original_runs = []
                    for run in paragraph.runs:
                        original_runs.append({
                            'text': run.text,
                            'bold': run.bold,
                            'italic': run.italic,
                            'underline': run.underline,
                            'color': run.font.color.rgb if run.font.color else None,
                            'size': run.font.size,
                            'name': run.font.name
                        })
                    
                    # Clear the paragraph to rebuild it
                    paragraph.clear()
                    
                    # Process each original run
                    for run_data in original_runs:
                        text = run_data['text']
                        i = 0
                        
                        while i < len(text):
                            # Check if the current character is an Arabic numeral
                            if text[i].isdigit():
                                # Collect the full number
                                num_start = i
                                while i < len(text) and text[i].isdigit():
                                    i += 1
                                number = text[num_start:i]
                                
                                # Add ornate parentheses around the number (U+FD3E & U+FD3F)
                                new_run = paragraph.add_run(f"﴿{number}﴾")
                                
                                # Apply original formatting to the new run
                                new_run.bold = run_data['bold']
                                new_run.italic = run_data['italic']
                                new_run.underline = run_data['underline']
                                if run_data['color']:
                                    new_run.font.color.rgb = run_data['color']
                                if run_data['size']:
                                    new_run.font.size = run_data['size']
                                if run_data['name']:
                                    new_run.font.name = run_data['name']
                                # Set RTL for the new run
                                new_run.font.rtl = True
                            else:
                                # Add the character as is with original formatting
                                new_run = paragraph.add_run(text[i])
                                
                                # Apply original formatting to the new run
                                new_run.bold = run_data['bold']
                                new_run.italic = run_data['italic']
                                new_run.underline = run_data['underline']
                                if run_data['color']:
                                    new_run.font.color.rgb = run_data['color']
                                if run_data['size']:
                                    new_run.font.size = run_data['size']
                                if run_data['name']:
                                    new_run.font.name = run_data['name']
                                # Set RTL for the new run
                                new_run.font.rtl = True
                                
                                i += 1
        
        # Save the modified document
        doc.save(target_path)
        
        print(f"Created new RTL document at: {target_path}")
        return target_path
    
    except Exception as e:
        print(f"Error processing DOCX file: {e}")
        return ""

def convert_docx_to_pdf(docx_path, pdf_path=None):
    """
    Convert a DOCX file to PDF using docx2pdf.
    
    Args:
        docx_path (str): Path to the DOCX file to convert
        pdf_path (str, optional): Path where the PDF will be saved.
                                If not provided, will use the same name with .pdf extension.
    
    Returns:
        str: Path to the created PDF file, or empty string if operation failed
    """
    try:
        # Validate source file exists
        if not os.path.exists(docx_path):
            print(f"Error: DOCX file not found at {docx_path}")
            return ""
        
        # Create PDF path if not provided
        if not pdf_path:
            pdf_path = os.path.splitext(docx_path)[0] + ".pdf"
        
        # Convert DOCX to PDF
        print(f"Converting {os.path.basename(docx_path)} to PDF...")
        convert(docx_path, pdf_path)
        
        print(f"PDF created successfully: {pdf_path}")
        return pdf_path
    
    except Exception as e:
        print(f"Error converting DOCX to PDF: {e}")
        return ""


if __name__ == "__main__":
    # Get the absolute path to the script's directory
    script_dir = os.path.dirname(os.path.abspath(__file__))
    
    # Check if file path was provided as argument
    if len(sys.argv) > 1:
        docx_path = sys.argv[1]
    else:
        # Default to the same file used in view_docx_content.py
        docx_path = os.path.join(script_dir, "pages", "602.docx")
    
    # Create new version
    new_path = center_docx_text(docx_path)
    
    if new_path:
        print(f"Successfully created new document: {new_path}")
    else:
        print("Failed to create new document.")
    
    # Convert new DOCX to PDF
    pdf_path = convert_docx_to_pdf(new_path)
    
    if pdf_path:
        print(f"Successfully created PDF: {pdf_path}")
    else:
        print("Failed to create PDF.")
