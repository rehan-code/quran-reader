import os
import docx
import sys

def get_docx_content(docx_path):
    """
    Open a DOCX file and return its text content.
    Returns a string containing all the text from the document.
    """
    try:
        # Open the DOCX file
        doc = docx.Document(docx_path)
        
        # Collect all text from paragraphs
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        
        # Join all paragraphs with newlines
        return "\n".join(full_text)
    
    except Exception as e:
        print(f"Error opening or reading file: {e}")
        return ""


def print_docx_details(docx_path):
    """
    Open a DOCX file and print detailed information about its contents,
    including Unicode characters and formatting runs.
    """
    try:
        # Open the DOCX file
        doc = docx.Document(docx_path)
        
        print(f"\nContents of {os.path.basename(docx_path)}:\n")
        print("-" * 60)
        
        # Process each paragraph
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip():
                print(f"Paragraph {i+1}:")
                
                # Print the raw text
                print(f"Raw text: {para.text}")
                
                # Print Unicode code points to help identify special characters
                print("Unicode code points:")
                for char in para.text:
                    print(f"  '{char}': U+{ord(char):04X}")
                
                # Print each run (formatting segment) separately
                print("Runs (formatting segments):")
                for j, run in enumerate(para.runs):
                    print(f"  Run {j+1}: '{run.text}'")
                
                print("-" * 40)
        
        print("-" * 60)
        return True
    
    except Exception as e:
        print(f"Error opening or reading file: {e}")
        return False

if __name__ == "__main__":
    # Get the absolute path to the script's directory
    script_dir = os.path.dirname(os.path.abspath(__file__))
    
    # Path to the 602.docx file
    docx_path = os.path.join(script_dir, "pages", "602.docx")
    
    if not os.path.exists(docx_path):
        print(f"Error: File not found at {docx_path}")
        sys.exit(1)
    
    # Get and print the text content of the file
    text_content = get_docx_content(docx_path)
    print(f"\nText content of {os.path.basename(docx_path)}:\n")
    print("-" * 60)
    print(text_content)
    print("-" * 60)
    
    # Optionally, print detailed information about the file
    # Uncomment the line below to see detailed Unicode information
    # print_docx_details(docx_path)
